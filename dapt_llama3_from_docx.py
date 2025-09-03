import argparse
import json
import os
import random
import re
from dataclasses import dataclass
from typing import List, Dict, Optional

from tqdm import tqdm

#   pip install python-docx transformers datasets accelerate peft bitsandbytes
try:
    from docx import Document as DocxDocument
except Exception as e:
    DocxDocument = None

import datasets
from datasets import load_dataset

import torch
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, Trainer, TrainingArguments,
    DataCollatorForLanguageModeling, BitsAndBytesConfig,   # ← 추가
)

try:
    from peft import LoraConfig, get_peft_model, prepare_model_for_kbit_training
except Exception:
    LoraConfig = None
    get_peft_model = None
    prepare_model_for_kbit_training = None

def read_docx_text(path: str) -> str:
    """Extract text from a .docx file (paragraphs + tables)."""
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed. Please: pip install python-docx")
    doc = DocxDocument(path)
    blocks: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            blocks.append(txt)

    # tables (flatten cells row-wise)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                blocks.append(" \t ".join(row_text))

    text = "\n".join(blocks)
    # Mild cleanup: collapse excessive spaces/newlines, normalize bullets
    text = re.sub(r"\u2022|\u25CF|\u25A0|\u2219", "-", text)  # bullets → dashes
    text = re.sub(r"\s+", " ", text)  # collapse whitespace
    text = re.sub(r"(\s*\-\s*){3,}", " - ", text)  # too many dashes → one dash
    return text.strip()

# Summarize subcommand (extractive, TextRank + MMR)

def split_sentences(text: str) -> List[str]:
    # 간단 문장 분리기: . ! ? 및 CJK 마침표 기준
    t = re.sub(r"[\r\t]", " ", text)
    t = re.sub(r"\s*-\s+", ". ", t)  # 불릿을 문장 경계로 보조 처리
    parts = re.split(r"(?<=[\.\!\?\u3002\uFF61])\s+", t)
    if len(parts) <= 1:
        parts = re.split(r"\s{2,}|;\s+", t)
    sents, seen = [], set()
    for s in parts:
        s = s.strip()
        if len(s) < 5: 
            continue
        if s in seen:
            continue
        seen.add(s)
        sents.append(s)
    return sents

def _build_tfidf(sentences: List[str]):
    # 문장 수가 너무 적으면 TF-IDF/텍스트랭크 대신 바로 폴백을 쓰게끔 신호 반환
    n = len(sentences)
    if n < 5:
        return None, None
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
    except Exception:
        return None, None

    # 작은 n에서 max_df 0.9가 min_df(1)보다 작게 매핑되는 문제 방지
    # n이 작을수록 max_df를 1.0로 열어 충돌을 피움
    max_df = 0.9 if n >= 20 else 1.0   # n<20이면 모든 문서 허용
    min_df = 1                         # 최소 1 문서에 등장

    vec = TfidfVectorizer(max_df=max_df, min_df=min_df, ngram_range=(1, 2))
    try:
        X = vec.fit_transform(sentences)
    except ValueError:
        # 드물게 어휘가 모두 필터링될 수도 있음 → 폴백
        return None, None
    return vec, X


def _cosine_sim_matrix(X):
    try:
        from sklearn.metrics.pairwise import cosine_similarity
    except Exception:
        return None
    return cosine_similarity(X, dense_output=False)

def _textrank_scores(sim_mat, d=0.85, max_iter=50, tol=1e-6):
    import numpy as np
    n = sim_mat.shape[0]
    if n == 0:
        return []
    row_sums = np.asarray(sim_mat.sum(axis=1)).ravel()
    P = sim_mat.copy()
    for i in range(n):
        if row_sums[i] > 0:
            P[i] = P[i] / row_sums[i]
    r = np.ones(n) / n
    for _ in range(max_iter):
        r_new = (1-d)/n + d * P.T.dot(r)
        if np.linalg.norm(r_new - r, 1) < tol:
            r = r_new
            break
        r = r_new
    return r.tolist()

def _mmr_select(sentences: List[str], scores: List[float], k: int, sim_mat=None, lambda_=0.7):
    n = len(sentences)
    if k >= n:
        return list(range(n))
    selected, candidate = [], set(range(n))
    while len(selected) < k and candidate:
        best_i, best_val = None, -1e9
        for i in list(candidate):
            rel = scores[i]
            div = 0.0
            if selected and sim_mat is not None:
                div = max(sim_mat[i, j] if i != j else 0.0 for j in selected)
            val = lambda_ * rel - (1 - lambda_) * div
            if val > best_val:
                best_val, best_i = val, i
        selected.append(best_i)
        candidate.remove(best_i)
    return selected

def summarize_text(text: str, target_ratio: float, max_sentences: int) -> str:
    sents = split_sentences(text)
    if not sents:
        return text
    k = max(1, min(max_sentences, int(len(sents) * target_ratio)))

    vec, X = _build_tfidf(sents)
    if vec is None or X is None:
        # 문장 수가 적거나 TF-IDF 구축 실패 → 간단 폴백
        return " ".join(sents[:k]).strip()

    sim = _cosine_sim_matrix(X)
    if sim is None:
        return " ".join(sents[:k]).strip()

    scores = _textrank_scores(sim)
    idx = _mmr_select(sents, scores, k=k, sim_mat=sim, lambda_=0.7)
    idx = sorted(idx)
    return " ".join(sents[i] for i in idx).strip()


def cmd_summarize(args):
    files = iter_docx_files(args.input_dir)
    if not files:
        raise FileNotFoundError(f"No .docx files found under: {args.input_dir}")
    recs = []
    for p in tqdm(files, desc="Summarizing .docx"):
        try:
            text = read_docx_text(p)
        except Exception as e:
            print(f"[WARN] Failed to parse {p}: {e}")
            continue
        if len(text) < args.min_chars:
            continue
        summ = summarize_text(text, target_ratio=args.target_ratio, max_sentences=args.max_sentences)
        # 요약이 너무 빈약하면 원문 유지(안전장치)
        if not summ or len(summ) < max(200, int(len(text) * 0.05)):
            summ = text
        recs.append({
            "text": summ,
            "source_path": os.path.relpath(p, args.input_dir),
            "summary": True,
        })
    write_jsonl(recs, args.output_jsonl)
    print(f"Wrote {len(recs)} summarized records to {args.output_jsonl}")



def iter_docx_files(input_dir: str) -> List[str]:
    paths = []
    for root, _, files in os.walk(input_dir):
        for f in files:
            if f.lower().endswith(".docx"):
                paths.append(os.path.join(root, f))
    paths.sort()
    return paths


def write_jsonl(records: List[Dict], out_path: str):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        for r in records:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")


# Preprocess subcommand

def cmd_preprocess(args):
    random.seed(args.seed)
    files = iter_docx_files(args.input_dir)
    if not files:
        raise FileNotFoundError(f"No .docx files found under: {args.input_dir}")

    records: List[Dict[str, str]] = []
    for p in tqdm(files, desc="Extracting .docx"):
        try:
            text = read_docx_text(p)
        except Exception as e:
            print(f"[WARN] Failed to parse {p}: {e}")
            continue
        if len(text) < args.min_chars:
            continue
        # One document per JSONL line under key 'text'
        records.append({"text": text, "source_path": os.path.relpath(p, args.input_dir)})

    if not records:
        raise RuntimeError("No documents survived preprocessing; try lowering --min_chars.")

    write_jsonl(records, args.output_jsonl)
    print(f"Wrote {len(records)} records to {args.output_jsonl}")


# Dataset tokenization & packing

@dataclass
class PackingConfig:
    block_size: Optional[int] = None  # If None → use tokenizer.model_max_length or 4096 fallback
    add_eos_between_docs: bool = True


def load_jsonl_as_dataset(train_jsonl: str, eval_ratio: float = 0.05, seed: int = 42):
    # 'all' 같은 예약 split명을 쓰지 말고, 단일 split로 로드한 뒤 나눕니다.
    # 결과: {"train": ..., "validation": ...}
    ds = load_dataset("json", data_files=train_jsonl, split="train")
    # 무작위 셔플 + train/val 분할
    ds = ds.shuffle(seed=seed)
    split = ds.train_test_split(test_size=max(1, int(len(ds) * eval_ratio)), seed=seed)
    # train_test_split은 test 비율/개수로 나누는데, 이름을 깔끔히 바꿔줍니다.
    return datasets.DatasetDict({
        "train": split["train"],
        "validation": split["test"],
    })



def build_tokenize_and_group_fn(tokenizer: AutoTokenizer, packing: PackingConfig):
    eos_id = tokenizer.eos_token_id
    if eos_id is None:
        # Common for LLaMA tokenizers: set pad to eos
        tokenizer.pad_token = tokenizer.eos_token
        eos_id = tokenizer.eos_token_id

    def tokenize(example):
        # Insert EOS at the end of each doc (optional)
        text = example["text"]
        if packing.add_eos_between_docs and tokenizer.eos_token:
            text = text + tokenizer.eos_token
        return tokenizer(text, add_special_tokens=False)

    # Determine block size
    if packing.block_size is None:
        # Use tokenizer.model_max_length when reasonable; cap to 4096 as a safe default
        try:
            max_len = tokenizer.model_max_length
            if isinstance(max_len, int) and max_len < 100000:  # not "int(1e30)"
                block_size = min(4096, max_len)
            else:
                block_size = 4096
        except Exception:
            block_size = 4096
    else:
        block_size = packing.block_size

    def group_texts(examples):
        # Concatenate and split into block_size chunks
        concatenated = []
        for ids in examples["input_ids"]:
            concatenated.extend(ids)
        # Drop the last partial chunk to follow typical DAPT convention
        total_length = (len(concatenated) // block_size) * block_size
        concatenated = concatenated[:total_length]
        result = {
            "input_ids": [concatenated[i : i + block_size] for i in range(0, total_length, block_size)]
        }
        result["attention_mask"] = [[1] * block_size for _ in range(len(result["input_ids"]))]
        # Labels for causal LM are the same as input_ids
        result["labels"] = [x.copy() for x in result["input_ids"]]
        return result

    return tokenize, group_texts, block_size


# Training subcommand

def maybe_prepare_lora(model, use_lora: bool, use_qlora: bool, lora_r: int, lora_alpha: int, lora_dropout: float):
    if not use_lora:
        return model
    if LoraConfig is None:
        raise RuntimeError("peft is not installed. Please: pip install peft")

    if use_qlora and prepare_model_for_kbit_training is not None:
        model = prepare_model_for_kbit_training(model)

    lora_cfg = LoraConfig(
        r=lora_r,
        lora_alpha=lora_alpha,
        lora_dropout=lora_dropout,
        bias="none",
        task_type="CAUSAL_LM",
        target_modules=[
            "q_proj", "k_proj", "v_proj", "o_proj",
            "gate_proj", "up_proj", "down_proj",
        ],
    )
    model = get_peft_model(model, lora_cfg)
    # Print trainable params
    try:
        model.print_trainable_parameters()
    except Exception:
        pass
    return model


def cmd_train(args):
    # Tokenizer
    tokenizer = AutoTokenizer.from_pretrained(args.model_id, use_fast=True)
    if tokenizer.pad_token is None and tokenizer.eos_token is not None:
        tokenizer.pad_token = tokenizer.eos_token

    # Model loading (BitsAndBytesConfig + 안전한 attention backend)
    load_kwargs = dict(
        device_map="auto",
        torch_dtype=torch.bfloat16 if args.bf16 else torch.float16,
        attn_implementation=getattr(args, "attn_impl", "eager"),  # torch<2.1이면 eager 권장
    )

    qconf = None
    if args.use_qlora:
        qconf = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_use_double_quant=True,
            bnb_4bit_compute_dtype=torch.bfloat16 if args.bf16 else torch.float16,
        )
    elif getattr(args, "use_int8", False):
        qconf = BitsAndBytesConfig(load_in_8bit=True)

    if qconf is not None:
        load_kwargs["quantization_config"] = qconf

    model = AutoModelForCausalLM.from_pretrained(args.model_id, **load_kwargs)


    if args.gradient_checkpointing:
        model.gradient_checkpointing_enable()

    # Prepare dataset
    ds = load_jsonl_as_dataset(args.train_jsonl, eval_ratio=args.eval_ratio, seed=args.seed)

    packing_cfg = PackingConfig(block_size=args.block_size if args.block_size > 0 else None,
                                add_eos_between_docs=not args.no_add_eos)
    tokenize_fn, group_fn, used_block = build_tokenize_and_group_fn(tokenizer, packing_cfg)

    tokenized = ds.map(tokenize_fn, batched=False, remove_columns=ds["train"].column_names)

    if args.packing:
        tokenized = tokenized.map(
            group_fn,
            batched=True,
            batch_size=args.group_batch_size,
        )

    collator = DataCollatorForLanguageModeling(tokenizer=tokenizer, mlm=False)

    os.makedirs(args.output_dir, exist_ok=True)

    training_args = TrainingArguments(
        output_dir=args.output_dir,
        num_train_epochs=args.num_train_epochs,
        per_device_train_batch_size=args.per_device_train_batch_size,
        per_device_eval_batch_size=args.per_device_eval_batch_size,
        gradient_accumulation_steps=args.gradient_accumulation_steps,
        learning_rate=args.learning_rate,
        weight_decay=args.weight_decay,
        warmup_ratio=args.warmup_ratio,
        lr_scheduler_type=args.lr_scheduler_type,
        logging_steps=args.logging_steps,
        save_steps=args.save_steps,
        eval_strategy="steps",
        eval_steps=args.eval_steps,
        save_total_limit=args.save_total_limit,
        bf16=args.bf16,
        fp16=not args.bf16,
        dataloader_num_workers=args.num_workers,
        gradient_checkpointing=args.gradient_checkpointing,
        optim="adamw_torch",
        report_to=["none"],
        prediction_loss_only=True,
    )

    # Apply (Q)LoRA if requested
    model = maybe_prepare_lora(
        model,
        use_lora=args.use_lora,
        use_qlora=args.use_qlora,
        lora_r=args.lora_r,
        lora_alpha=args.lora_alpha,
        lora_dropout=args.lora_dropout,
    )

    trainer = Trainer(
        model=model,
        args=training_args,
        train_dataset=tokenized["train"],
        eval_dataset=tokenized["validation"],
        data_collator=collator,
    )

    trainer.train()
    eval_out = trainer.evaluate()
    try:
        ppl = float(torch.exp(torch.tensor(eval_out["eval_loss"])) )
        print(f"Perplexity: {ppl:.2f}")
    except Exception:
        pass

    # Save final adapter or full model
    if args.use_lora:
        model.save_pretrained(os.path.join(args.output_dir, "adapter"))
        tokenizer.save_pretrained(args.output_dir)
        print(f"Saved LoRA adapter to {os.path.join(args.output_dir, 'adapter')}\n"
              f"To merge adapters into the base model later, use PEFT utilities or inference with model+adapter.")
    else:
        trainer.save_model(args.output_dir)
        tokenizer.save_pretrained(args.output_dir)
        print(f"Saved full model to {args.output_dir}")


# CLI

def build_arg_parser():
    p = argparse.ArgumentParser(description="DOCX → JSONL → DAPT for Llama‑3 Instruct")
    sub = p.add_subparsers(dest="command", required=True)

    sp = sub.add_parser("preprocess", help="Extract text from .docx to JSONL")
    sp.add_argument("--input_dir", type=str, required=True, help="Folder containing .docx files")
    sp.add_argument("--output_jsonl", type=str, required=True, help="Path to write JSONL")
    sp.add_argument("--min_chars", type=int, default=800, help="Drop docs shorter than this")
    sp.add_argument("--seed", type=int, default=42)
    sp.set_defaults(func=cmd_preprocess)

    ss = sub.add_parser("summarize", help="Summarize .docx to compressed JSONL (extractive)")
    ss.add_argument("--input_dir", type=str, required=True)
    ss.add_argument("--output_jsonl", type=str, required=True)
    ss.add_argument("--target_ratio", type=float, default=0.25, help="Fraction of sentences to keep")
    ss.add_argument("--max_sentences", type=int, default=40)
    ss.add_argument("--min_chars", type=int, default=500)
    ss.set_defaults(func=cmd_summarize)

    st = sub.add_parser("train", help="Run continued pre‑training (causal LM)")
    st.add_argument("--use_int8", action="store_true")
    st.add_argument("--attn_impl", type=str, default="eager", choices=["eager","sdpa","flash_attention_2"],
                    help="Attention backend. 'eager'는 가장 호환성이 높음")
    st.add_argument("--model_id", type=str, required=True,
                    help="HF model id or local path (e.g., meta-llama/Meta-Llama-3-8B-Instruct)")
    st.add_argument("--train_jsonl", type=str, required=True, help="JSONL from preprocess (has 'text')")
    st.add_argument("--output_dir", type=str, required=True)

    # Data & packing
    st.add_argument("--eval_ratio", type=float, default=0.05)
    st.add_argument("--block_size", type=int, default=-1, help="Token block size; -1 → auto (≈4096)")
    st.add_argument("--group_batch_size", type=int, default=1000,
                    help="How many examples to group per map() call when packing")
    st.add_argument("--no_add_eos", action="store_true", help="Do not append EOS between docs")
    st.add_argument("--packing", action="store_true", help="Pack sequences up to block_size")

    # Training hyperparams
    st.add_argument("--num_train_epochs", type=int, default=10)
    st.add_argument("--per_device_train_batch_size", type=int, default=2)
    st.add_argument("--per_device_eval_batch_size", type=int, default=2)
    st.add_argument("--gradient_accumulation_steps", type=int, default=8)
    st.add_argument("--learning_rate", type=float, default=2e-5)
    st.add_argument("--weight_decay", type=float, default=0.05)
    st.add_argument("--warmup_ratio", type=float, default=0.1)
    st.add_argument("--lr_scheduler_type", type=str, default="cosine")
    st.add_argument("--logging_steps", type=int, default=2)
    st.add_argument("--save_steps", type=int, default=200)
    st.add_argument("--eval_steps", type=int, default=2)
    st.add_argument("--save_total_limit", type=int, default=2)
    st.add_argument("--num_workers", type=int, default=2)

    # Mixed precision & memory
    st.add_argument("--bf16", action="store_true")
    st.add_argument("--gradient_checkpointing", action="store_true")

    # (Q)LoRA options
    st.add_argument("--use_lora", action="store_true")
    st.add_argument("--use_qlora", action="store_true")
    st.add_argument("--lora_r", type=int, default=16)
    st.add_argument("--lora_alpha", type=int, default=32)
    st.add_argument("--lora_dropout", type=float, default=0.05)

    st.add_argument("--seed", type=int, default=42)

    st.set_defaults(func=cmd_train)
    return p


def main():
    parser = build_arg_parser()
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
